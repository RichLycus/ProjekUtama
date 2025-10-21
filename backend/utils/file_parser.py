"""
File Parser Utilities
Parse various file formats: PDF, DOCX, TXT
"""

import os
from pathlib import Path
from typing import Dict, Any, Optional
import PyPDF2
from docx import Document
from PIL import Image
import io


class FileParser:
    """Parse different file formats and extract content"""
    
    @staticmethod
    def parse_pdf(file_path: str) -> Dict[str, Any]:
        """
        Parse PDF file and extract text content
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dictionary with parsed content and metadata
        """
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from all pages
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append({
                            'page': page_num,
                            'content': page_text.strip()
                        })
                
                # Combine all text
                full_text = '\n\n'.join([p['content'] for p in text_content])
                
                return {
                    'success': True,
                    'content': full_text,
                    'pages': len(pdf_reader.pages),
                    'page_contents': text_content,
                    'metadata': {
                        'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                        'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                        'subject': pdf_reader.metadata.get('/Subject', '') if pdf_reader.metadata else '',
                    }
                }
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'content': ''
            }
    
    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, Any]:
        """
        Parse DOCX file and extract text content
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with parsed content and metadata
        """
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables if any
            tables_content = []
            for table_idx, table in enumerate(doc.tables, start=1):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_content.append({
                    'table_number': table_idx,
                    'rows': table_data
                })
            
            # Combine content
            full_text = '\n\n'.join(paragraphs)
            
            # Add tables to text
            if tables_content:
                full_text += '\n\n--- Tables ---\n\n'
                for table in tables_content:
                    full_text += f"Table {table['table_number']}:\n"
                    for row in table['rows']:
                        full_text += ' | '.join(row) + '\n'
                    full_text += '\n'
            
            return {
                'success': True,
                'content': full_text,
                'paragraphs': len(paragraphs),
                'tables': len(tables_content),
                'metadata': {
                    'paragraphs_count': len(paragraphs),
                    'tables_count': len(tables_content)
                }
            }
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'content': ''
            }
    
    @staticmethod
    def parse_txt(file_path: str) -> Dict[str, Any]:
        """
        Parse TXT file and extract content
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Dictionary with parsed content and metadata
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            lines = content.split('\n')
            non_empty_lines = [line for line in lines if line.strip()]
            
            return {
                'success': True,
                'content': content,
                'metadata': {
                    'lines': len(lines),
                    'non_empty_lines': len(non_empty_lines),
                    'characters': len(content)
                }
            }
        except Exception as e:
            # Try different encodings
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                
                return {
                    'success': True,
                    'content': content,
                    'metadata': {
                        'encoding': 'latin-1',
                        'characters': len(content)
                    }
                }
            except:
                return {
                    'success': False,
                    'error': str(e),
                    'content': ''
                }
    
    @staticmethod
    def process_image(file_path: str) -> Dict[str, Any]:
        """
        Process image file and extract metadata
        
        Args:
            file_path: Path to image file
            
        Returns:
            Dictionary with image metadata
        """
        try:
            with Image.open(file_path) as img:
                # Get image info
                metadata = {
                    'format': img.format,
                    'mode': img.mode,
                    'size': img.size,
                    'width': img.width,
                    'height': img.height,
                }
                
                # Check if image has EXIF data
                if hasattr(img, '_getexif') and img._getexif():
                    metadata['has_exif'] = True
                else:
                    metadata['has_exif'] = False
                
                return {
                    'success': True,
                    'metadata': metadata
                }
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }
    
    @staticmethod
    def parse_file(file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Parse file based on its type
        
        Args:
            file_path: Path to file
            file_type: File extension (pdf, docx, txt, etc.)
            
        Returns:
            Parsed content and metadata
        """
        file_type = file_type.lower().replace('.', '')
        
        if file_type == 'pdf':
            return FileParser.parse_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            return FileParser.parse_docx(file_path)
        elif file_type in ['txt', 'md', 'csv']:
            return FileParser.parse_txt(file_path)
        elif file_type in ['png', 'jpg', 'jpeg', 'webp', 'gif']:
            return FileParser.process_image(file_path)
        else:
            return {
                'success': False,
                'error': f'Unsupported file type: {file_type}'
            }
